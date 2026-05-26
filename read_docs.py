import os, glob
from docx import Document

doc_dir = r"c:\Users\Lenovo\Desktop\Code2Change\Documentation"
output_file = r"c:\Users\Lenovo\Desktop\Code2Change\all_docs_text.txt"

files = sorted(glob.glob(os.path.join(doc_dir, "*.docx")))

with open(output_file, "w", encoding="utf-8") as out:
    for f in files:
        basename = os.path.basename(f)
        out.write(f"\n{'='*80}\n")
        out.write(f"FILE: {basename}\n")
        out.write(f"{'='*80}\n\n")
        try:
            doc = Document(f)
            for para in doc.paragraphs:
                out.write(para.text + "\n")
            # Also extract tables
            for i, table in enumerate(doc.tables):
                out.write(f"\n[TABLE {i+1}]\n")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    out.write(" | ".join(cells) + "\n")
                out.write("\n")
        except Exception as e:
            out.write(f"ERROR reading file: {e}\n")

print(f"Done. Output written to {output_file}")
